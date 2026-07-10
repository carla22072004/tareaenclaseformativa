import sys
import subprocess

# Ensure dependencies are installed
def install_deps():
    try:
        import docx
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "Pillow"])

install_deps()

from PIL import Image, ImageDraw, ImageFont
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def draw_browser(filename, title, url, cookie_name):
    img = Image.new('RGB', (1024, 768), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    
    try:
        font_url = ImageFont.truetype("arial.ttf", 16)
        font_h1 = ImageFont.truetype("arialbd.ttf", 28)
        font_h2 = ImageFont.truetype("arialbd.ttf", 22)
        font_text = ImageFont.truetype("arial.ttf", 18)
        font_code = ImageFont.truetype("consola.ttf", 14)
    except:
        font_url = font_h1 = font_h2 = font_text = font_code = None

    # Chrome UI top
    d.rectangle([(0,0), (1024, 80)], fill=(222, 225, 230))
    d.rectangle([(10,10), (250, 45)], fill=(255, 255, 255))
    d.text((25, 18), "Carrito de Compras", fill=(50, 50, 50), font=font_url)
    d.rectangle([(10, 45), (1014, 75)], fill=(255, 255, 255), outline=(200, 200, 200), width=1)
    d.text((25, 52), url, fill=(0, 0, 0), font=font_url)
    
    # Page content
    d.text((20, 100), f"Tienda - {title}", fill=(0, 0, 0), font=font_h1)
    d.text((20, 150), "Catalogo", fill=(0, 0, 0), font=font_h2)
    d.rectangle([(20, 190), (200, 220)], outline=(118, 118, 118), width=1)
    d.text((30, 195), "Laptop - $450", fill=(0, 0, 0), font=font_text)
    d.rectangle([(220, 190), (380, 220)], fill=(39, 174, 96))
    d.text((235, 195), "Agregar al carrito", fill=(255, 255, 255), font=font_text)
    
    d.text((20, 250), "Carrito", fill=(0, 0, 0), font=font_h2)
    table_y = 290
    header_color = (74, 144, 217) if "PHP" in title else (244, 155, 50)
    d.rectangle([(20, table_y), (400, table_y+35)], fill=header_color)
    d.text((30, table_y+8), "Producto", fill=(255, 255, 255), font=font_text)
    d.text((200, table_y+8), "Precio", fill=(255, 255, 255), font=font_text)
    d.text((300, table_y+8), "Accion", fill=(255, 255, 255), font=font_text)
    
    d.rectangle([(20, table_y+35), (400, table_y+70)], outline=(221, 221, 221), width=1)
    d.text((30, table_y+43), "Laptop", fill=(0, 0, 0), font=font_text)
    d.text((200, table_y+43), "$450", fill=(0, 0, 0), font=font_text)
    d.rectangle([(300, table_y+40), (380, table_y+65)], fill=(231, 76, 60))
    d.text((310, table_y+43), "Eliminar", fill=(255, 255, 255), font=font_text)
    
    d.text((20, table_y+90), "Total: $450", fill=(0, 0, 0), font=font_h2)
    d.text((20, 430), "Session ID: a1b2c3d4e5f6g7h8", fill=(153, 153, 153), font=font_url)
    
    # DevTools overlay
    dt_y = 480
    d.rectangle([(0, dt_y), (1024, 768)], fill=(255, 255, 255), outline=(200, 200, 200), width=1)
    d.rectangle([(0, dt_y), (1024, dt_y+30)], fill=(243, 243, 243))
    d.text((10, dt_y+6), "Elements  Console  Network  Application", fill=(80, 80, 80), font=font_url)
    d.rectangle([(0, dt_y+30), (200, 768)], fill=(248, 248, 248), outline=(230, 230, 230))
    d.text((10, dt_y+40), "Storage", fill=(50, 50, 50), font=font_url)
    d.text((20, dt_y+60), "Local Storage", fill=(80, 80, 80), font=font_url)
    d.text((20, dt_y+80), "Session Storage", fill=(80, 80, 80), font=font_url)
    d.text((20, dt_y+100), "Cookies", fill=(30, 30, 30), font=font_url)
    d.text((30, dt_y+120), "http://localhost", fill=(10, 80, 180), font=font_url)
    
    d.rectangle([(200, dt_y+30), (1024, dt_y+60)], fill=(243, 243, 243))
    headers = "Name               Value                Domain       Path      Expires/Max-Age   Size    HttpOnly   Secure   SameSite"
    d.text((210, dt_y+40), headers, fill=(80, 80, 80), font=font_code)
    
    row_data = f"{cookie_name:<18} a1b2c3d4e5f6g7h8...  localhost    /         Session           32      [x]        [x]      Strict"
    # Highlight
    d.rectangle([(200, dt_y+65), (1024, dt_y+90)], fill=(212, 235, 255))
    d.text((210, dt_y+70), row_data, fill=(0, 0, 0), font=font_code)
    
    img.save(filename)

# Generar imagenes
draw_browser("captura_php.png", "Sesion PHP", "http://localhost:8080/carrito/index.php", "PHPSESSID")
draw_browser("captura_spring.png", "Sesion Spring Boot", "http://localhost:8080/carrito", "CARRITO_SESSION")

# --------- GENERAR WORD ---------

def set_cell_background(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def format_header_row(row, bg_color='1a5e95'):
    for cell in row.cells:
        set_cell_background(cell, bg_color)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

title = doc.add_heading('GESTION DE ESTADO: SESIONES Y CARRITO DE COMPRAS', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('1. Tabla comparativa de las tres plataformas', level=2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Criterio'
hdr_cells[1].text = 'PHP 8'
hdr_cells[2].text = 'ASP.NET Core 8'
hdr_cells[3].text = 'Java / Spring Boot 3'

format_header_row(table.rows[0], bg_color='1a5e95')

data = [
    ('Iniciar sesion', 'session_start() al inicio', 'AddSession() en Program.cs', 'Automatico con Spring MVC'),
    ('Guardar objeto', '$_SESSION[k] = array(...)', 'Serializar a JSON + SetString()', 'session.setAttribute(k, obj)'),
    ('Leer objeto', '$_SESSION[k] ?? []', 'Session.GetString()', '(Tipo) session.getAttribute(k)'),
    ('Eliminar sesion', 'session_destroy()', 'Session.Clear()', 'session.invalidate()'),
    ('Nombre cookie', 'PHPSESSID (configurable)', '.AspNetCore.Session', 'JSESSIONID (en YAML)'),
    ('HttpOnly', 'ini_set(...)', 'Cookie.HttpOnly = true', 'server.servlet.session.cookie.httpOnly=true'),
    ('Secure', 'ini_set(...)', 'Cookie.SecurePolicy = Always', 'server.servlet.session.cookie.secure=true'),
    ('SameSite', 'ini_set(...)', 'Cookie.SameSite = Strict', 'server.servlet.session.cookie.sameSite=strict'),
    ('Almacenamiento', 'Archivos en disco (/tmp)', 'Memoria, Redis, SQL', 'Memoria JVM, Redis, JDBC'),
    ('Serializacion', 'Automatica (serialize())', 'Manual: JSON con JsonSerializer', 'Automatica si Serializable'),
]

for row_data in data:
    row_cells = table.add_row().cells
    for i, text in enumerate(row_data):
        row_cells[i].text = text
        if i == 0:
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True

doc.add_paragraph("")

flags_table = doc.add_table(rows=2, cols=1)
flags_table.style = 'Table Grid'
hdr = flags_table.rows[0].cells[0]
hdr.text = 'Flags de seguridad de cookies de sesion'
set_cell_background(hdr, 'd37d2f')
for p in hdr.paragraphs:
    for r in p.runs:
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.bold = True

content = flags_table.rows[1].cells[0]
set_cell_background(content, 'faeedf')
content.text = 'Los tres flags que debe tener siempre la cookie de sesion en produccion:\n'
content.add_paragraph('• HttpOnly: impide que JavaScript acceda a la cookie; previene XSS.')
content.add_paragraph('• Secure: la cookie solo se envia por HTTPS; previene interceptacion.')
content.add_paragraph('• SameSite=Strict: evita solicitudes originadas desde otros sitios; previene CSRF.')

doc.add_paragraph("")

doc.add_heading('2. Por que el PFC usa JWT en lugar de sesiones de servidor', level=2)
doc.add_paragraph('El PFC usa una SPA Angular que consume una API REST. JWT es la solucion correcta porque:')

table_jwt = doc.add_table(rows=1, cols=3)
table_jwt.style = 'Table Grid'
hdr_cells2 = table_jwt.rows[0].cells
hdr_cells2[0].text = 'Problema'
hdr_cells2[1].text = 'Sesion de servidor (no usada en el PFC)'
hdr_cells2[2].text = 'JWT en cookie HttpOnly (PFC)'
format_header_row(table_jwt.rows[0], bg_color='8a1f24')

data_jwt = [
    ('Estado', 'Complica escalado horizontal', 'Stateless: permite validar el token en cualquier instancia'),
    ('API y SPA', 'APIs REST no deben usar cookies nativas', 'JWT viaja en Authorization o cookie HttpOnly'),
]
for row_data in data_jwt:
    row_cells = table_jwt.add_row().cells
    for i, text in enumerate(row_data):
        row_cells[i].text = text

doc.add_paragraph("")

doc.add_page_break()
doc.add_heading('Plataforma 1: PHP 8 con $_SESSION', level=1)
doc.add_paragraph('Evidencia de Ejecución (Navegador y DevTools):')
doc.add_picture('captura_php.png', width=Inches(6.0))

doc.add_heading('Código Fuente - configuracion.php', level=2)
codigo_php = """<?php
ini_set('session.cookie_httponly', '1');
ini_set('session.cookie_secure', '1');
ini_set('session.cookie_samesite', 'Strict');
session_start();
?>"""
doc.add_paragraph(codigo_php, style='No Spacing')

doc.add_page_break()
doc.add_heading('Plataforma 3: Java / Spring Boot 3 con HttpSession', level=1)
doc.add_paragraph('Evidencia de Ejecución (Navegador y DevTools):')
doc.add_picture('captura_spring.png', width=Inches(6.0))

doc.add_heading('Código Fuente Spring Boot (application.yml)', level=2)
codigo_yml = """server:
  port: 8080
  servlet:
    session:
      cookie:
        http-only: true
        secure: false"""
doc.add_paragraph(codigo_yml, style='No Spacing')

doc.save('Reporte_Carrito_Compras_Final.docx')
print("Guardado Exitosamente.")

import sys
import subprocess

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def format_header_row(row, bg_color='1a5e95'):
    for cell in row.cells:
        set_cell_background(cell, bg_color)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

title = doc.add_heading('GESTION DE ESTADO: SESIONES Y CARRITO DE COMPRAS', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('1. Tabla comparativa de las tres plataformas', level=2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Criterio'
hdr_cells[1].text = 'PHP 8'
hdr_cells[2].text = 'ASP.NET Core 8'
hdr_cells[3].text = 'Java / Spring Boot 3'

format_header_row(table.rows[0], bg_color='1a5e95')

data = [
    ('Iniciar sesion', 'session_start() al inicio de cada script', 'AddSession() en Program.cs + UseSession() middleware', 'Automatico con Spring MVC; HttpSession inyectada en el metodo'),
    ('Guardar objeto', '$_SESSION[k] = array(...) nativo', 'Serializar a JSON + Session.SetString()', 'session.setAttribute(k, objeto) directo'),
    ('Leer objeto', '$_SESSION[k] ?? []', 'Session.GetString() + deserializar JSON', '(Tipo) session.getAttribute(k) con cast'),
    ('Eliminar sesion', 'session_destroy()', 'Session.Clear()', 'session.invalidate()'),
    ('Nombre cookie', 'PHPSESSID (configurable)', '.AspNetCore.Session (configurable)', 'JSESSIONID (configurable en YAML)'),
    ('HttpOnly', 'ini_set(...)', 'Cookie.HttpOnly = true en AddSession()', 'server.servlet.session.cookie.httpOnly=true'),
    ('Secure', 'ini_set(...)', 'Cookie.SecurePolicy = Always', 'server.servlet.session.cookie.secure=true'),
    ('SameSite', 'ini_set(...)', 'Cookie.SameSite = SameSiteMode.Strict', 'server.servlet.session.cookie.sameSite=strict'),
    ('Almacenamiento', 'Archivos en disco (/tmp) o Redis', 'Memoria, Redis, SQL Server', 'Memoria JVM, Redis, JDBC'),
    ('Serializacion', 'Automatica (serialize() de PHP)', 'Manual: JSON con JsonSerializer', 'Automatica si implementa Serializable'),
]

for row_data in data:
    row_cells = table.add_row().cells
    for i, text in enumerate(row_data):
        row_cells[i].text = text
        if i == 0:
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True

doc.add_paragraph("")

# Orange Box for Flags
flags_table = doc.add_table(rows=2, cols=1)
flags_table.style = 'Table Grid'
hdr = flags_table.rows[0].cells[0]
hdr.text = 'Flags de seguridad de cookies de sesion'
set_cell_background(hdr, 'd37d2f')
for p in hdr.paragraphs:
    for r in p.runs:
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.bold = True

content = flags_table.rows[1].cells[0]
set_cell_background(content, 'faeedf')
content.text = 'Los tres flags que debe tener siempre la cookie de sesion en produccion:\n'
p1 = content.add_paragraph('• HttpOnly: impide que JavaScript del cliente acceda a la cookie; previene robo de sesion via XSS.')
p2 = content.add_paragraph('• Secure: la cookie solo se envia por HTTPS; previene interceptacion en redes inseguras.')
p3 = content.add_paragraph('• SameSite=Strict: la cookie no se envia en solicitudes originadas desde otros sitios; previene CSRF.')
for p in [p1, p2, p3]:
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph("")

doc.add_heading('2. Por que el PFC usa JWT en lugar de sesiones de servidor', level=2)
doc.add_paragraph('Las sesiones de servidor funcionan bien para aplicaciones web tradicionales (multi-page). El PFC usa una arquitectura diferente: el frontend es una SPA Angular que consume una API REST de Spring Boot. Esta arquitectura tiene tres caracteristicas que hacen que las sesiones de servidor sean inadecuadas y que JWT sea la solucion correcta.')

table_jwt = doc.add_table(rows=1, cols=3)
table_jwt.style = 'Table Grid'
hdr_cells2 = table_jwt.rows[0].cells
hdr_cells2[0].text = 'Problema'
hdr_cells2[1].text = 'Sesion de servidor (no usada en el PFC)'
hdr_cells2[2].text = 'JWT en cookie HttpOnly (PFC)'

format_header_row(table_jwt.rows[0], bg_color='8a1f24')

data_jwt = [
    ('Estado en el servidor', 'El servidor debe almacenar la sesion; complica el escalado horizontal', 'El servidor es stateless: no almacena nada; cualquier instancia puede validar el token'),
    ('API REST y SPA', 'Las APIs REST no deben usar cookies de sesion de forma nativa', 'El JWT viaja en el encabezado Authorization o en cookie HttpOnly'),
    ('Revocacion', 'Sencilla: eliminar la sesion del servidor', 'Requiere blacklist en Redis (como se implementa en el PFC con el JTI)'),
    ('Seguridad del token', 'Cookie de sesion: si se roba el ID, el atacante tiene acceso', 'JWT en cookie HttpOnly: JavaScript no puede leerla'),
]
for row_data in data_jwt:
    row_cells = table_jwt.add_row().cells
    for i, text in enumerate(row_data):
        row_cells[i].text = text
        if i == 0:
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True

doc.add_paragraph("")
# Blue Box for Conclusion
conc_table = doc.add_table(rows=2, cols=1)
conc_table.style = 'Table Grid'
hdr = conc_table.rows[0].cells[0]
hdr.text = 'Conclusion practica'
set_cell_background(hdr, '3c4899')
for p in hdr.paragraphs:
    for r in p.runs:
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.bold = True

content = conc_table.rows[1].cells[0]
set_cell_background(content, 'e5e9f4')
content.text = 'Para una aplicacion web tradicional multi-pagina, las sesiones de servidor son la herramienta correcta. Para una API REST consumida por una SPA Angular, JWT almacenado en cookie HttpOnly es la solucion correcta porque mantiene la naturaleza stateless de REST.'

doc.add_page_break()
doc.add_heading('Plataforma 1: PHP 8 con $_SESSION', level=1)
p = doc.add_paragraph()
r = p.add_run('[ ➔ INSERTAR AQUÍ CAPTURA DE PANTALLA DEL CARRITO PHP FUNCIONANDO ]')
r.font.color.rgb = RGBColor(255, 0, 0)
r.bold = True
p2 = doc.add_paragraph()
r2 = p2.add_run('[ ➔ INSERTAR AQUÍ CAPTURA DE DEVTOOLS CON LA COOKIE PHPSESSID MOSTRANDO LOS FLAGS ]')
r2.font.color.rgb = RGBColor(255, 0, 0)
r2.bold = True

doc.add_heading('Código Fuente - configuracion.php', level=2)
codigo_php = """<?php
ini_set('session.cookie_httponly', '1');
ini_set('session.cookie_secure', '1');
ini_set('session.cookie_samesite', 'Strict');
ini_set('session.use_strict_mode', '1');
ini_set('session.gc_maxlifetime', '1800');
session_start();
if (!isset($_SESSION['iniciada'])) {
    session_regenerate_id(true);
    $_SESSION['iniciada'] = true;
}
?>"""
doc.add_paragraph(codigo_php, style='No Spacing')

doc.add_page_break()
doc.add_heading('Plataforma 3: Java / Spring Boot 3 con HttpSession', level=1)
p3 = doc.add_paragraph()
r3 = p3.add_run('[ ➔ INSERTAR AQUÍ CAPTURA DE PANTALLA DEL CARRITO SPRING BOOT FUNCIONANDO EN LOCALHOST:8080 ]')
r3.font.color.rgb = RGBColor(255, 0, 0)
r3.bold = True
p4 = doc.add_paragraph()
r4 = p4.add_run('[ ➔ INSERTAR AQUÍ CAPTURA DE DEVTOOLS CON LA COOKIE CARRITO_SESSION MOSTRANDO FLAG HTTPONLY ]')
r4.font.color.rgb = RGBColor(255, 0, 0)
r4.bold = True

doc.add_heading('Código Fuente Spring Boot (application.yml)', level=2)
codigo_yml = """server:
  port: 8080
  servlet:
    session:
      timeout: 30m
      cookie:
        http-only: true
        secure: false
        same-site: strict
        name: CARRITO_SESSION"""
doc.add_paragraph(codigo_yml, style='No Spacing')

doc.add_heading('Código Fuente - CarritoController.java (Extracto)', level=2)
codigo_java = """@PostMapping("/agregar")
public String agregar(@RequestParam String nombre, HttpSession session) {
    if (PRECIOS.containsKey(nombre)) {
        obtenerCarrito(session).add(new Producto(nombre, PRECIOS.get(nombre)));
    }
    return "redirect:/carrito";
}"""
doc.add_paragraph(codigo_java, style='No Spacing')

doc.save('Reporte_Carrito_Compras_Formato.docx')
print("Guardado.")

import sys
import subprocess

# Instalar python-docx si no existe
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Titulo principal
title = doc.add_heading('Reporte de Práctica: Sesiones y Carrito de Compras', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Análisis Teórico y Comparativa
doc.add_heading('1. Análisis Teórico y Comparativa', level=1)
doc.add_heading('Tabla Comparativa de las tres plataformas', level=2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Criterio'
hdr_cells[1].text = 'PHP 8'
hdr_cells[2].text = 'ASP.NET Core 8'
hdr_cells[3].text = 'Java / Spring Boot 3'

data = [
    ('Iniciar sesion', 'session_start() al inicio', 'AddSession() en Program.cs', 'HttpSession inyectada en el metodo'),
    ('Guardar objeto', '$_SESSION[k] = array(...)', 'Serializar a JSON + SetString()', 'session.setAttribute(k, objeto)'),
    ('Leer objeto', '$_SESSION[k] ?? []', 'GetString() + deserializar JSON', 'session.getAttribute(k) con cast'),
    ('Eliminar sesion', 'session_destroy()', 'Session.Clear()', 'session.invalidate()'),
    ('HttpOnly', 'ini_set(...)', 'Cookie.HttpOnly = true', 'server.servlet.session.cookie.httpOnly=true'),
    ('Secure', 'ini_set(...)', 'Cookie.SecurePolicy = Always', 'server.servlet.session.cookie.secure=true'),
    ('SameSite', 'ini_set(...)', 'Cookie.SameSite = Strict', 'server.servlet.session.cookie.sameSite=strict'),
]
for criterio, php, asp, java in data:
    row_cells = table.add_row().cells
    row_cells[0].text = criterio
    row_cells[1].text = php
    row_cells[2].text = asp
    row_cells[3].text = java

doc.add_heading('Flags de seguridad de cookies de sesion', level=2)
doc.add_paragraph('Para que una cookie de sesión se considere segura, debe tener:')
doc.add_paragraph('1. HttpOnly: impide que JavaScript acceda a la cookie (previene ataques XSS).')
doc.add_paragraph('2. Secure: la cookie solo viaja por HTTPS (previene interceptación en red o Man-in-the-Middle).')
doc.add_paragraph('3. SameSite=Strict: evita el envío de la cookie en solicitudes cruzadas (previene ataques CSRF).')

doc.add_heading('Justificación: JWT vs Sesiones de Servidor', level=2)
p = doc.add_paragraph('En el PFC se utiliza una arquitectura SPA (Angular) + API REST (Spring Boot). ')
p.add_run('Las sesiones de servidor violan el principio stateless de REST y complican el escalado horizontal. ')
p.add_run('Se emplea un JWT almacenado en una cookie HttpOnly porque: mantiene el backend stateless, permite escalar horizontalmente fácilmente al no depender de la memoria del servidor, y protege contra ataques XSS al ser inaccesible desde JavaScript.')

# 2. PHP Code and Photos
doc.add_page_break()
doc.add_heading('2. Plataforma 1: PHP 8 con $_SESSION', level=1)
doc.add_paragraph('A continuación, se presenta la evidencia de ejecución y el código fuente.')

doc.add_heading('Evidencia de Ejecución (Captura de pantalla)', level=2)
p = doc.add_paragraph()
r = p.add_run('[ ➔ INSERTAR AQUÍ CAPTURA DE PANTALLA DEL NAVEGADOR CON EL CARRITO PHP FUNCIONANDO ]')
r.font.color.rgb = RGBColor(255, 0, 0)
r.bold = True

p = doc.add_paragraph()
r = p.add_run('[ ➔ INSERTAR AQUÍ CAPTURA DE DEVTOOLS (F12) MOSTRANDO LA COOKIE PHPSESSID CON FLAGS HTTPONLY Y SECURE ]')
r.font.color.rgb = RGBColor(255, 0, 0)
r.bold = True

doc.add_heading('Código Fuente - configuracion.php', level=2)
codigo_php = '''<?php
ini_set('session.cookie_httponly', '1');
ini_set('session.cookie_secure', '1');
ini_set('session.cookie_samesite', 'Strict');
ini_set('session.use_strict_mode', '1');
ini_set('session.gc_maxlifetime', '1800');
session_start();
if (!isset($_SESSION['iniciada'])) {
    session_regenerate_id(true);
    $_SESSION['iniciada'] = true;
}
?>'''
doc.add_paragraph(codigo_php, style='No Spacing')

# 3. Java Spring Boot Code and Photos
doc.add_page_break()
doc.add_heading('3. Plataforma 3: Java / Spring Boot 3 con HttpSession', level=1)
doc.add_paragraph('A continuación, se presenta la evidencia de ejecución y fragmentos del código fuente de Spring Boot.')

doc.add_heading('Evidencia de Ejecución (Captura de pantalla)', level=2)
p = doc.add_paragraph()
r = p.add_run('[ ➔ INSERTAR AQUÍ CAPTURA DE PANTALLA DEL NAVEGADOR CON EL CARRITO SPRING BOOT (LOCALHOST:8080) ]')
r.font.color.rgb = RGBColor(255, 0, 0)
r.bold = True

p = doc.add_paragraph()
r = p.add_run('[ ➔ INSERTAR AQUÍ CAPTURA DE DEVTOOLS (F12) MOSTRANDO LA COOKIE CARRITO_SESSION CON FLAG HTTPONLY ]')
r.font.color.rgb = RGBColor(255, 0, 0)
r.bold = True

doc.add_heading('Código Fuente - application.yml', level=2)
codigo_yml = '''server:
  port: 8080
  servlet:
    session:
      timeout: 30m
      cookie:
        http-only: true
        secure: false
        same-site: strict
        name: CARRITO_SESSION'''
doc.add_paragraph(codigo_yml, style='No Spacing')

doc.add_heading('Código Fuente - CarritoController.java (Extracto de Agregar)', level=2)
codigo_java = '''@PostMapping("/agregar")
public String agregar(@RequestParam String nombre, HttpSession session) {
    if (PRECIOS.containsKey(nombre)) {
        obtenerCarrito(session).add(new Producto(nombre, PRECIOS.get(nombre)));
    }
    return "redirect:/carrito";
}'''
doc.add_paragraph(codigo_java, style='No Spacing')

doc.save('Reporte_Carrito_Compras.docx')
print("Documento guardado exitosamente.")
